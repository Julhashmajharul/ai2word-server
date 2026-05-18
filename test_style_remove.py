import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document()
h = doc.add_heading('Test Heading with Math', level=1)
h.add_run(' some text')

pPr = h._element.get_or_add_pPr()
pStyle = pPr.find(qn('w:pStyle'))
if pStyle is not None:
    pPr.remove(pStyle) # remove style completely

doc.save('test_style_remove.docx')
print("File saved")
