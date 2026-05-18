"""
DOCX processing utilities — extracted from the original app.py.
All Word document styling, font splitting, math alignment, heading flattening,
Bengali font injection, table border, horizontal line VML, and page border logic.
"""
import io
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Helper – extract primary font name from CSS font stack
# ---------------------------------------------------------------------------
def _first_font_name(font_family_str: str) -> str:
    first = font_family_str.split(",")[0].strip().strip("'\"")
    return first if first else "Times New Roman"


# ---------------------------------------------------------------------------
# Font helpers
# ---------------------------------------------------------------------------
def force_bengali_english_font(run, ascii_font="Times New Roman", cs_font="bangla"):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    for old_fonts in rPr.findall(qn('w:rFonts')):
        rPr.remove(old_fonts)
    new_fonts = OxmlElement('w:rFonts')
    new_fonts.set(qn('w:ascii'),    ascii_font)
    new_fonts.set(qn('w:hAnsi'),    ascii_font)
    new_fonts.set(qn('w:eastAsia'), ascii_font)
    new_fonts.set(qn('w:cs'),       cs_font)
    rPr.append(new_fonts)


def safe_apply_fonts_preserve_math(paragraph, ascii_font="Times New Roman", bengali_font="bangla"):
    for run in list(paragraph.runs):
        if not run.text.strip():
            continue
        chunks = re.split(r'([\u0980-\u09FF\u0964\u0965]+)', run.text)
        parent = run._element.getparent()
        for chunk in chunks:
            if not chunk:
                continue
            new_run = paragraph.add_run(chunk)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size: new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb: new_run.font.color.rgb = run.font.color.rgb
            if re.search(r'[\u0980-\u09FF\u0964\u0965]', chunk):
                new_run.font.name = bengali_font
            else:
                new_run.font.name = ascii_font
            rPr = new_run._element.get_or_add_rPr()
            for rFonts in rPr.findall(qn('w:rFonts')):
                rPr.remove(rFonts)
            new_fonts = OxmlElement('w:rFonts')
            if re.search(r'[\u0980-\u09FF\u0964\u0965]', chunk):
                new_fonts.set(qn('w:ascii'), bengali_font)
                new_fonts.set(qn('w:hAnsi'), bengali_font)
                new_fonts.set(qn('w:cs'), bengali_font)
            else:
                new_fonts.set(qn('w:ascii'), ascii_font)
                new_fonts.set(qn('w:hAnsi'), ascii_font)
                new_fonts.set(qn('w:cs'), ascii_font)
            rPr.append(new_fonts)
            parent.insert(parent.index(run._element), new_run._element)
        parent.remove(run._element)


def fix_math_bengali_fonts(doc, bengali_font):
    for m_r in doc._element.body.findall('.//' + qn('m:r')):
        t_node = m_r.find(qn('m:t'))
        if t_node is not None and t_node.text and re.search(r'[\u0980-\u09FF\u0964\u0965]', t_node.text):
            m_rPr = m_r.find(qn('m:rPr'))
            if m_rPr is None:
                m_rPr = OxmlElement('m:rPr')
                m_r.insert(0, m_rPr)
            if m_rPr.find(qn('m:nor')) is None:
                m_rPr.append(OxmlElement('m:nor'))
            w_rPr = m_r.find(qn('w:rPr'))
            if w_rPr is None:
                w_rPr = OxmlElement('w:rPr')
                m_r.insert(0, w_rPr)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), bengali_font)
            rFonts.set(qn('w:hAnsi'), bengali_font)
            rFonts.set(qn('w:cs'), bengali_font)
            for old_f in w_rPr.findall(qn('w:rFonts')):
                w_rPr.remove(old_f)
            w_rPr.append(rFonts)


# ---------------------------------------------------------------------------
# Math alignment
# ---------------------------------------------------------------------------
def set_math_alignment(doc, align_val):
    omml_align = "left" if align_val == "left" else "centerGroup"
    for math_para in doc._element.body.findall('.//' + qn('m:oMathPara')):
        pr = math_para.find(qn('m:oMathParaPr'))
        if pr is None:
            pr = OxmlElement('m:oMathParaPr')
            math_para.insert(0, pr)
        jc = pr.find(qn('m:jc'))
        if jc is None:
            jc = OxmlElement('m:jc')
            pr.append(jc)
        jc.set(qn('m:val'), omml_align)
        parent = math_para.getparent()
        while parent is not None and parent.tag != qn('w:p'):
            parent = parent.getparent()
        if parent is not None:
            p = Paragraph(parent, parent.getparent())
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align_val == "left" else WD_ALIGN_PARAGRAPH.CENTER
        for eqArr in math_para.findall('.//' + qn('m:eqArr')):
            eqArrPr = eqArr.find(qn('m:eqArrPr'))
            if eqArrPr is None:
                eqArrPr = OxmlElement('m:eqArrPr')
                eqArr.insert(0, eqArrPr)
            baseJc = eqArrPr.find(qn('m:baseJc'))
            if baseJc is None:
                baseJc = OxmlElement('m:baseJc')
                eqArrPr.append(baseJc)
            baseJc.set(qn('m:val'), omml_align)
    if align_val == "left":
        for math_inline in doc._element.body.findall('.//' + qn('m:oMath')):
            parent = math_inline.getparent()
            while parent is not None and parent.tag != qn('w:p'):
                parent = parent.getparent()
            if parent is not None:
                p = Paragraph(parent, parent.getparent())
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ---------------------------------------------------------------------------
# Column layout
# ---------------------------------------------------------------------------
def set_document_columns(section, count):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sect_pr.append(cols)
    cols.set(qn('w:num'), str(count))
    if count > 1:
        cols.set(qn('w:sep'), '1')
        cols.set(qn('w:space'), '432')
    else:
        cols.attrib.pop(qn('w:sep'), None)
        cols.attrib.pop(qn('w:space'), None)


# ---------------------------------------------------------------------------
# Page borders
# ---------------------------------------------------------------------------
_CSS_TO_WORD_BORDER = {
    "solid": "single", "dashed": "dashed", "dotted": "dotted",
    "double": "double", "none": "none",
}

def _apply_page_borders(section, border_style, border_width_px, border_color):
    word_val = _CSS_TO_WORD_BORDER.get(border_style.lower(), "single")
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(existing)
    if word_val == "none":
        return
    sz = max(2, int(border_width_px * 6))
    color_hex = border_color.lstrip("#").upper().zfill(6)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border_el = OxmlElement(f"w:{side}")
        border_el.set(qn("w:val"), word_val)
        border_el.set(qn("w:sz"), str(sz))
        border_el.set(qn("w:space"), "24")
        border_el.set(qn("w:color"), color_hex)
        pg_borders.append(border_el)
    sect_pr.append(pg_borders)


# ---------------------------------------------------------------------------
# Table borders
# ---------------------------------------------------------------------------
def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)
    else:
        tblPr = tblPr[0]
    tblBorders = tblPr.xpath('w:tblBorders')
    if not tblBorders:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else:
        tblBorders = tblBorders[0]
        tblBorders.clear()
    for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
        border = OxmlElement(border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)


# ---------------------------------------------------------------------------
# Main docx styling pipeline
# ---------------------------------------------------------------------------
def apply_docx_settings(docx_path: str, settings: dict) -> io.BytesIO:
    """Open Pandoc-generated DOCX and apply all user settings."""
    doc = Document(docx_path)

    # Margins
    top    = Inches(float(settings.get("margin_top_in",    0.7)))
    bottom = Inches(float(settings.get("margin_bottom_in", 0.7)))
    left   = Inches(float(settings.get("margin_left_in",   0.5)))
    right  = Inches(float(settings.get("margin_right_in",  0.5)))

    border_style = settings.get("border_style", "none")
    border_width = int(settings.get("border_width_px", 2))
    border_color = settings.get("border_color", "#000000")
    column_count = int(settings.get("column_count", 1))

    for section in doc.sections:
        section.top_margin    = top
        section.bottom_margin = bottom
        section.left_margin   = left
        section.right_margin  = right
        _apply_page_borders(section, border_style, border_width, border_color)
        set_document_columns(section, column_count)

    # Font
    font_name = _first_font_name(settings.get("font_family", "Times New Roman"))
    font_size = Pt(float(settings.get("font_size_pt", 12)))
    cs_font_raw      = settings.get("cs_font", "").strip()
    cs_font_effective = cs_font_raw if cs_font_raw else font_name

    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = font_size

    # Paragraph spacing
    space_before = Pt(float(settings.get("space_before_pt", 0)))
    space_after  = Pt(float(settings.get("space_after_pt",  0)))
    normal_style.paragraph_format.space_before = space_before
    normal_style.paragraph_format.space_after  = space_after

    class _FakeRun:
        def __init__(self, element): self._element = element
    force_bengali_english_font(_FakeRun(normal_style.element), font_name, cs_font_effective)

    # HR settings
    hr_enabled = settings.get("hr_enabled", True)
    hr_width  = settings.get("hr_width", "100%")
    hr_height = settings.get("hr_height", 1)
    hr_color  = settings.get("hr_color", "#000000")
    hr_align  = settings.get("hr_align", "center")

    # PASS 1 – Body paragraphs
    for paragraph in doc.paragraphs:
        if 'SINGLE_LINE_PLACEHOLDER' in paragraph.text:
            if hr_enabled:
                paragraph.clear()
                hr_xml = f"""
                <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                     xmlns:v="urn:schemas-microsoft-com:vml"
                     xmlns:o="urn:schemas-microsoft-com:office:office">
                    <w:pict>
                        <v:rect style="width:{hr_width};height:{hr_height}pt" o:hr="t" o:hrstd="t" o:hrnoshade="t" o:hralign="{hr_align}" fillcolor="{hr_color}" stroked="f"/>
                    </w:pict>
                </w:r>
                """
                hr_run = parse_xml(hr_xml.strip())
                paragraph._p.append(hr_run)
                if hr_align == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif hr_align == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = paragraph._element
                p.getparent().remove(p)
            continue

        if paragraph.text.strip():
            safe_apply_fonts_preserve_math(paragraph, font_name, cs_font_effective)
        paragraph.paragraph_format.space_before = space_before
        paragraph.paragraph_format.space_after  = space_after

    # Table cells
    for table in doc.tables:
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        if 'XYZZYBRXYZZY' in run.text:
                            parts = run.text.split('XYZZYBRXYZZY')
                            run.text = parts[0]
                            for part in parts[1:]:
                                run.add_break()
                                run.add_text(part)
                    if 'SINGLE_LINE_PLACEHOLDER' in paragraph.text:
                        if hr_enabled:
                            paragraph.clear()
                            hr_xml = f"""
                            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                                 xmlns:v="urn:schemas-microsoft-com:vml"
                                 xmlns:o="urn:schemas-microsoft-com:office:office">
                                <w:pict>
                                    <v:rect style="width:{hr_width};height:{hr_height}pt" o:hr="t" o:hrstd="t" o:hrnoshade="t" o:hralign="{hr_align}" fillcolor="{hr_color}" stroked="f"/>
                                </w:pict>
                            </w:r>
                            """
                            hr_run = parse_xml(hr_xml.strip())
                            paragraph._p.append(hr_run)
                            if hr_align == 'left': paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif hr_align == 'right': paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p = paragraph._element
                            p.getparent().remove(p)
                        continue
                    if paragraph.text.strip():
                        safe_apply_fonts_preserve_math(paragraph, font_name, cs_font_effective)

    # PASS 2 – Headings
    _HEADING_SIZES = {1: Pt(16), 2: Pt(14)}
    _DEFAULT_HEADING_SIZE = Pt(13)
    header_color_hex = settings.get("header_color", "#003366").lstrip("#")
    try:
        header_rgb = RGBColor.from_string(header_color_hex)
    except Exception:
        header_rgb = RGBColor(0x00, 0x33, 0x66)

    for para in doc.paragraphs:
        style_name = para.style.name
        if not style_name.startswith("Heading"):
            continue
        try:
            depth = int(style_name.split()[-1])
        except (ValueError, IndexError):
            depth = 3
        heading_pt = _HEADING_SIZES.get(depth, _DEFAULT_HEADING_SIZE)
        pPr = para._element.get_or_add_pPr()
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            pPr.remove(pStyle)
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), '9')
        for run in para.runs:
            run.font.bold      = True
            run.font.size      = heading_pt
            run.font.color.rgb = header_rgb

    # PASS 3 – Bengali math fonts
    fix_math_bengali_fonts(doc, cs_font_effective)

    # PASS 4 – Math alignment
    math_align = settings.get("math_align", "center")
    set_math_alignment(doc, math_align)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
