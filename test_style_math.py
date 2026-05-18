import docx
from docx.oxml import parse_xml

# 1. Create a document with a heading that contains math
doc = docx.Document()
p = doc.add_paragraph()
p.style = doc.styles["Heading 1"]

# Add some math
math_xml = '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t>x=y</m:t></m:r></m:oMath></m:oMathPara>'
p._p.append(parse_xml(math_xml))
doc.save('test_math_heading.docx')

# 2. Re-open, change style to Normal, and save again
doc2 = docx.Document('test_math_heading.docx')
p2 = doc2.paragraphs[0]
p2.style = doc2.styles["Normal"]
doc2.save('test_math_normal.docx')
print("Done creating test files")
