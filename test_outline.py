import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document()
doc.add_heading('Test Heading', level=1)
doc.add_paragraph('Some body text.')

p = doc.paragraphs[0]

# Add outlineLvl 9 to pPr
pPr = p._element.get_or_add_pPr()
outlineLvl = pPr.find(qn('w:outlineLvl'))
if outlineLvl is None:
    outlineLvl = OxmlElement('w:outlineLvl')
    pPr.append(outlineLvl)
outlineLvl.set(qn('w:val'), '9')

doc.save('test_outline.docx')
print("Saved test_outline.docx")
