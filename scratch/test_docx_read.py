from docx import Document

doc = Document("test.docx")
table = doc.tables[0]
cell = table.cell(1, 0)
with open("out.xml", "w", encoding="utf-8") as f:
    for p in cell.paragraphs:
        f.write("Text: " + p.text + "\n")
        f.write("XML: " + p._p.xml + "\n")
