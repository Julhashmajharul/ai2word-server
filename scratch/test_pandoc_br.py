import subprocess
from docx import Document

md = """| A |
|---|
| 1. $VC = \\frac{x}{n}$ XYZZYBRXYZZY XYZZYBRXYZZY 2. $VC = s - x$ |"""

with open("test.md", "w", encoding="utf-8") as f:
    f.write(md)

subprocess.run([
    "pandoc", "test.md", "-o", "test.docx",
    "--from=markdown+tex_math_dollars+tex_math_single_backslash"
])

doc = Document("test.docx")
table = doc.tables[0]
for p in table.cell(1,0).paragraphs:
    for r in p.runs:
        print(f"Run text: {repr(r.text)}")
