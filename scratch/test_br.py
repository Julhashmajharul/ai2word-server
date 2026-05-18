import docx

doc = docx.Document()
p = doc.add_paragraph()
r = p.add_run("Hello@@BR@@World")

text_parts = r.text.split('@@BR@@')
r.text = text_parts[0]
for part in text_parts[1:]:
    r.add_break()
    r.add_text(part)

doc.save("test_br.docx")
print("Saved")
