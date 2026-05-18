import io
import os
import re
import subprocess
import tempfile

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

# This pattern will now see through multiple lines inside $$ blocks
display_math_pattern = re.compile(r'\$\$(.*?)\$\$', re.DOTALL)

app = Flask(__name__)
CORS(app)  # Allow requests from your frontend domain


# ---------------------------------------------------------------------------
# Helper – apply user settings to an existing DOCX with python-docx
# ---------------------------------------------------------------------------

def _first_font_name(font_family_str: str) -> str:
    """
    Extract the primary font name from a CSS-style font stack.
    e.g. "'Times New Roman', 'Kalpurush', serif"  →  "Times New Roman"
    """
    first = font_family_str.split(",")[0].strip().strip("'\"")
    return first if first else "Times New Roman"


def force_bengali_english_font(run, ascii_font="Times New Roman", cs_font="bangla"):
    """
    Brutally replace any existing w:rFonts on a run with a fresh element that
    explicitly sets all four font slots:
      w:ascii / w:hAnsi  → ascii_font  (Latin text — e.g. Times New Roman)
      w:eastAsia         → ascii_font  (prevents Word doing its own CJK fallback)
      w:cs               → cs_font     (complex-script / Indic — e.g. bangla)
    Uses manual XML find/insert to work on raw lxml elements (style.element, etc.)
    which do NOT have the python-docx get_or_add_rPr() method.
    """
    # Manual XML handling — safe for raw elements that lack get_or_add_rPr()
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)

    for old_fonts in rPr.findall(qn('w:rFonts')):
        rPr.remove(old_fonts)

    new_fonts = OxmlElement('w:rFonts')
    new_fonts.set(qn('w:ascii'),    ascii_font)
    new_fonts.set(qn('w:hAnsi'),    ascii_font)
    new_fonts.set(qn('w:eastAsia'), ascii_font)
    new_fonts.set(qn('w:cs'),       cs_font)
    rPr.append(new_fonts)


def safe_apply_fonts_preserve_math(paragraph, ascii_font="Times New Roman", bengali_font="bangla"):
    # Iterate over existing text runs ONLY. This completely ignores and safely preserves <m:oMath>!
    for run in list(paragraph.runs):
        if not run.text.strip():
            continue

        chunks = re.split(r'([\u0980-\u09FF\u0964\u0965]+)', run.text)
        parent = run._element.getparent()
        
        for chunk in chunks:
            if not chunk: 
                continue
            
            # Create a new run temporarily at the end of the paragraph
            new_run = paragraph.add_run(chunk)
            
            # Copy basic formatting from the original run
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size: new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb: new_run.font.color.rgb = run.font.color.rgb
            
            if re.search(r'[\u0980-\u09FF\u0964\u0965]', chunk):
                new_run.font.name = bengali_font
            else:
                new_run.font.name = ascii_font
            
            # Apply strict font definitions via OXML
            rPr = new_run._element.get_or_add_rPr()
            for rFonts in rPr.findall(qn('w:rFonts')):
                rPr.remove(rFonts)
                
            new_fonts = OxmlElement('w:rFonts')
            if re.search(r'[\u0980-\u09FF\u0964\u0965]', chunk):
                new_fonts.set(qn('w:ascii'), bengali_font)
                new_fonts.set(qn('w:hAnsi'), bengali_font)
                new_fonts.set(qn('w:cs'), bengali_font)
            else:
                new_fonts.set(qn('w:ascii'), ascii_font)
                new_fonts.set(qn('w:hAnsi'), ascii_font)
                new_fonts.set(qn('w:cs'), ascii_font)
            rPr.append(new_fonts)
            
            # IMPORTANT: Move the newly created run's XML element exactly BEFORE the original run
            parent.insert(parent.index(run._element), new_run._element)
        
        # Finally, delete the original run from the XML tree without touching any math nodes
        parent.remove(run._element)


def fix_math_bengali_fonts(doc, bengali_font):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re

    # Search entire document body instead of just doc.paragraphs to catch math in tables/lists
    for m_r in doc._element.body.findall('.//' + qn('m:r')):
        t_node = m_r.find(qn('m:t'))
        if t_node is not None and t_node.text and re.search(r'[\u0980-\u09FF\u0964\u0965]', t_node.text):

            # m:rPr (Math Run Properties) হ্যান্ডলিং
            m_rPr = m_r.find(qn('m:rPr'))
            if m_rPr is None:
                m_rPr = OxmlElement('m:rPr')
                m_r.insert(0, m_rPr)

            # যুক্তবর্ণ ঠিক রাখতে m:nor (Normal Text) যোগ করা
            if m_rPr.find(qn('m:nor')) is None:
                m_rPr.append(OxmlElement('m:nor'))

            # w:rPr (Word Run Properties) হ্যান্ডলিং
            w_rPr = m_r.find(qn('w:rPr'))
            if w_rPr is None:
                w_rPr = OxmlElement('w:rPr')
                m_r.insert(0, w_rPr)

            # ফন্ট সেট করা
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), bengali_font)
            rFonts.set(qn('w:hAnsi'), bengali_font)
            rFonts.set(qn('w:cs'), bengali_font)

            for old_f in w_rPr.findall(qn('w:rFonts')):
                w_rPr.remove(old_f)
            w_rPr.append(rFonts)


from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph

def set_math_alignment(doc, align_val):
    """Set OMML equation alignment in the Word document.

    Word OMML uses 'm:jc' (justification) on 'm:oMathParaPr' to control
    display-math alignment. Values: 'centerGroup' (default) or 'left'.
    """
    omml_align = "left" if align_val == "left" else "centerGroup"
    p_align_val = "left" if align_val == "left" else "center"

    # 1. Handle Display Math (m:oMathPara)
    # Search entire document body instead of just doc.paragraphs to catch math in tables/lists
    for math_para in doc._element.body.findall('.//' + qn('m:oMathPara')):
        pr = math_para.find(qn('m:oMathParaPr'))
        if pr is None:
            pr = OxmlElement('m:oMathParaPr')
            math_para.insert(0, pr)

        jc = pr.find(qn('m:jc'))
        if jc is None:
            jc = OxmlElement('m:jc')
            pr.append(jc)

        jc.set(qn('m:val'), omml_align)

        # Force the parent paragraph to also be left-aligned to ensure no centering inheritance
        parent = math_para.getparent()
        while parent is not None and parent.tag != qn('w:p'):
            parent = parent.getparent()
        if parent is not None:
            p = Paragraph(parent, parent.getparent())
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align_val == "left" else WD_ALIGN_PARAGRAPH.CENTER

        # Force internal equation arrays (\begin{aligned}) to align left
        for eqArr in math_para.findall('.//' + qn('m:eqArr')):
            eqArrPr = eqArr.find(qn('m:eqArrPr'))
            if eqArrPr is None:
                eqArrPr = OxmlElement('m:eqArrPr')
                eqArr.insert(0, eqArrPr)
            baseJc = eqArrPr.find(qn('m:baseJc'))
            if baseJc is None:
                baseJc = OxmlElement('m:baseJc')
                eqArrPr.append(baseJc)
            # baseJc only supports specific values, left/right/center/centerGroup. We use the mapped omml_align.
            baseJc.set(qn('m:val'), omml_align)

    # 2. Handle Inline Math (m:oMath)
    # If a paragraph contains inline math and we want textbook style, force the paragraph to left-align.
    # This catches equations that Pandoc didn't wrap in m:oMathPara (e.g. lines with text like "=> $...$")
    if align_val == "left":
        for math_inline in doc._element.body.findall('.//' + qn('m:oMath')):
            parent = math_inline.getparent()
            while parent is not None and parent.tag != qn('w:p'):
                parent = parent.getparent()
            if parent is not None:
                p = Paragraph(parent, parent.getparent())
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_document_columns(section, count):
    """Set the number of columns for a Word document section via OXML.

    Word uses <w:cols w:num="N" w:sep="1" w:space="432"/> inside <w:sectPr>.
    w:space="432" = 432 twentieths-of-a-point = 0.3 inch column gap.
    w:sep="1" draws a vertical line between columns.
    """
    sect_pr = section._sectPr
    cols = sect_pr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sect_pr.append(cols)
    cols.set(qn('w:num'),   str(count))
    if count > 1:
        cols.set(qn('w:sep'),   '1')    # vertical divider between columns
        cols.set(qn('w:space'), '432')  # 0.3 inch gap
    else:
        # Single column — remove sep/space if they exist
        cols.attrib.pop(qn('w:sep'),   None)
        cols.attrib.pop(qn('w:space'), None)


# Map CSS border-style names → Word border val attribute names
_CSS_TO_WORD_BORDER = {
    "solid":  "single",
    "dashed": "dashed",
    "dotted": "dotted",
    "double": "double",
    "none":   "none",
}


def _apply_page_borders(section, border_style: str, border_width_px: int, border_color: str):
    """
    Inject a <w:pgBorders> element into the given section's <w:sectPr>.
    python-docx has no native API for page borders, so we build the OXML
    directly.

    border_style   : CSS value ("solid", "dashed", "dotted", "double", "none")
    border_width_px: integer pixels from the UI (we convert to eighth-points;
                     1 px ≈ 0.75 pt, 1 pt = 8 eighth-pts  →  factor ≈ 6)
    border_color   : hex string with or without leading "#" (e.g. "#000000")
    """
    word_val = _CSS_TO_WORD_BORDER.get(border_style.lower(), "single")

    # Remove any existing <w:pgBorders> to avoid duplicates
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(existing)

    if word_val == "none":
        return  # No border requested – we're done after removing old ones

    # Convert px → eighth-points (Word's unit for border size).
    # 1 px = 0.75 pt; 1 pt = 8 eighth-pts  →  1 px ≈ 6 eighth-pts
    sz = max(2, int(border_width_px * 6))   # minimum 2 (≈ 0.25 pt)

    # Normalise hex color: strip "#", ensure 6-char uppercase
    color_hex = border_color.lstrip("#").upper().zfill(6)

    # Build <w:pgBorders w:offsetFrom="page">
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")

    for side in ("top", "left", "bottom", "right"):
        border_el = OxmlElement(f"w:{side}")
        border_el.set(qn("w:val"),   word_val)
        border_el.set(qn("w:sz"),    str(sz))
        border_el.set(qn("w:space"), "24")   # 24 pt gap from page edge
        border_el.set(qn("w:color"), color_hex)
        pg_borders.append(border_el)

    # <w:pgBorders> must come after <w:pgMar> in sectPr – append at end
    sect_pr.append(pg_borders)

def set_table_borders(table):
    """Apply a 0.5pt solid black border to all cells in a table."""
    tblPr = table._element.xpath('w:tblPr')
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)
    else:
        tblPr = tblPr[0]
        
    tblBorders = tblPr.xpath('w:tblBorders')
    if not tblBorders:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else:
        tblBorders = tblBorders[0]
        tblBorders.clear()

    for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
        border = OxmlElement(border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)


def apply_docx_settings(docx_path: str, settings: dict):
    """
    Open the DOCX produced by Pandoc and apply:
      - Page margins (all four sides, in inches)
      - Page borders via OXML (python-docx has no native API for this)
      - Default body font family and size (via the 'Normal' style)
    Saves the file in-place.
    """
    doc = Document(docx_path)

    # --- Margins ---
    top    = Inches(float(settings.get("margin_top_in",    0.7)))
    bottom = Inches(float(settings.get("margin_bottom_in", 0.7)))
    left   = Inches(float(settings.get("margin_left_in",   0.5)))
    right  = Inches(float(settings.get("margin_right_in",  0.5)))

    # --- Page borders (OXML) ---
    border_style = settings.get("border_style", "none")
    border_width = int(settings.get("border_width_px", 2))
    border_color = settings.get("border_color", "#000000")
    column_count = int(settings.get("column_count", 1))

    for section in doc.sections:
        section.top_margin    = top
        section.bottom_margin = bottom
        section.left_margin   = left
        section.right_margin  = right

        # Apply page border via OXML
        _apply_page_borders(section, border_style, border_width, border_color)

        # Apply column layout
        set_document_columns(section, column_count)

    # --- Font family & size via the 'Normal' paragraph style ---
    font_name = _first_font_name(settings.get("font_family", "Times New Roman"))
    font_size = Pt(float(settings.get("font_size_pt", 12)))
    # cs_font: dedicated complex-script font (e.g. "bangla" for Bengali text).
    # If not provided or empty, falls back to font_name (single-font mode).
    cs_font_raw      = settings.get("cs_font", "").strip()
    cs_font_effective = cs_font_raw if cs_font_raw else font_name
    has_split_font   = bool(cs_font_raw)  # True when ascii != complex-script

    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = font_size

    # --- Paragraph spacing ---
    space_before = Pt(float(settings.get("space_before_pt", 0)))
    space_after  = Pt(float(settings.get("space_after_pt",  0)))
    normal_style.paragraph_format.space_before = space_before
    normal_style.paragraph_format.space_after  = space_after

    # Stamp rFonts on the Normal style via the helper so the document-wide
    # default also carries the correct complex-script font.
    class _FakeRun:
        """Thin wrapper so force_bengali_english_font() can accept a style element."""
        def __init__(self, element): self._element = element
    force_bengali_english_font(_FakeRun(normal_style.element), font_name, cs_font_effective)

    # --- Horizontal Line settings ---
    hr_enabled = settings.get("hr_enabled", True)
    hr_width  = settings.get("hr_width", "100%")
    hr_height = settings.get("hr_height", 1)
    hr_color  = settings.get("hr_color", "#000000")
    hr_align  = settings.get("hr_align", "center")

    # ---------------------------------------------------------------------------
    # PASS 1 – Regex-rebuild every paragraph and table cell.
    # Splits runs at Bengali Unicode boundaries and assigns font.name per chunk.
    # This is the only foolproof method — word's complex-script engine cannot
    # override an explicit run.font.name set by python-docx on a per-Run object.
    # ---------------------------------------------------------------------------

    # All body paragraphs
    for paragraph in doc.paragraphs:
        if 'SINGLE_LINE_PLACEHOLDER' in paragraph.text:
            if hr_enabled:
                paragraph.clear()
                
                # Dynamically build the VML shape using values from the frontend
                hr_xml = f"""
                <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                     xmlns:v="urn:schemas-microsoft-com:vml" 
                     xmlns:o="urn:schemas-microsoft-com:office:office">
                    <w:pict>
                        <v:rect style="width:{hr_width};height:{hr_height}pt" o:hr="t" o:hrstd="t" o:hrnoshade="t" o:hralign="{hr_align}" fillcolor="{hr_color}" stroked="f"/>
                    </w:pict>
                </w:r>
                """
                hr_run = parse_xml(hr_xml.strip())
                paragraph._p.append(hr_run)
                
                # Center the paragraph itself to ensure alignment works properly
                if hr_align == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif hr_align == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Completely remove the paragraph element so no blank space is left
                p = paragraph._element
                p.getparent().remove(p)
            
            continue # Skip the regex font rebuild for this line
            
        # 3. For all other normal text paragraphs, run the font splitting logic:
        if paragraph.text.strip():
            safe_apply_fonts_preserve_math(paragraph, font_name, cs_font_effective)

        # Always apply paragraph spacing to override Pandoc's defaults
        paragraph.paragraph_format.space_before = space_before
        paragraph.paragraph_format.space_after  = space_after

    # All paragraphs inside every table cell
    for table in doc.tables:
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Handle soft line breaks inside table cells
                    for run in paragraph.runs:
                        if 'XYZZYBRXYZZY' in run.text:
                            parts = run.text.split('XYZZYBRXYZZY')
                            run.text = parts[0]
                            for part in parts[1:]:
                                run.add_break()
                                run.add_text(part)

                    if 'SINGLE_LINE_PLACEHOLDER' in paragraph.text:
                        if hr_enabled:
                            paragraph.clear()
                            
                            hr_xml = f"""
                            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                                 xmlns:v="urn:schemas-microsoft-com:vml" 
                                 xmlns:o="urn:schemas-microsoft-com:office:office">
                                <w:pict>
                                    <v:rect style="width:{hr_width};height:{hr_height}pt" o:hr="t" o:hrstd="t" o:hrnoshade="t" o:hralign="{hr_align}" fillcolor="{hr_color}" stroked="f"/>
                                </w:pict>
                            </w:r>
                            """
                            hr_run = parse_xml(hr_xml.strip())
                            paragraph._p.append(hr_run)
                            
                            if hr_align == 'left':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif hr_align == 'right':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # Completely remove the paragraph element so no blank space is left
                            p = paragraph._element
                            p.getparent().remove(p)
                        continue
                            
                    if paragraph.text.strip():
                        safe_apply_fonts_preserve_math(paragraph, font_name, cs_font_effective)

    # ---------------------------------------------------------------------------
    # PASS 2 – Headings: flatten to Normal AFTER the body loop so the heading
    # bold/color/size cannot be overwritten by the body pass above.
    # ---------------------------------------------------------------------------
    _HEADING_SIZES = {1: Pt(16), 2: Pt(14)}
    _DEFAULT_HEADING_SIZE = Pt(13)

    # --- Header color ---
    header_color_hex = settings.get("header_color", "#003366").lstrip("#")
    try:
        header_rgb = RGBColor.from_string(header_color_hex)
    except Exception:
        header_rgb = RGBColor(0x00, 0x33, 0x66)

    # PASS 2 – Headings: flatten to Normal AFTER the body loop so heading
    # bold/color/size cannot be overwritten by PASS 1.
    for para in doc.paragraphs:
        style_name = para.style.name           # e.g. "Heading 1"
        if not style_name.startswith("Heading"):
            continue

        try:
            depth = int(style_name.split()[-1])
        except (ValueError, IndexError):
            depth = 3
        heading_pt = _HEADING_SIZES.get(depth, _DEFAULT_HEADING_SIZE)

        # Convert to Normal → removes outline level / collapsible triangle
        # WARNING: Changing style to Normal hides paragraphs containing <m:oMath> in some Word versions!
        # para.style = doc.styles["Normal"]

        # Instead, explicitly set the paragraph outline level to 9 (Body Text)
        # to remove the collapsible dropdown triangle without changing the style.
        pPr = para._element.get_or_add_pPr()
        
        # Remove the Heading style reference completely to avoid the collapsible triangle.
        # This acts exactly like 'Clear Formatting' in MS Word.
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            pPr.remove(pStyle)
            
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), '9')

        # Stamp every heading run — runs last so bold/color/size win
        # Font assignment was already done by rebuild_paragraph_with_fonts above.
        for run in para.runs:
            run.font.bold      = True
            run.font.size      = heading_pt
            run.font.color.rgb = header_rgb

    # ---------------------------------------------------------------------------
    # PASS 3 – Fix Bengali fonts inside OMML math equations.
    # The paragraph.runs loop above deliberately skips <m:oMath> elements.
    # This pass traverses the raw XML to stamp Bengali fonts on <m:r>/<m:t> nodes.
    # ---------------------------------------------------------------------------
    fix_math_bengali_fonts(doc, cs_font_effective)

    # ---------------------------------------------------------------------------
    # PASS 4 – Set math alignment (left / center) on all display math blocks.
    # ---------------------------------------------------------------------------
    math_align = settings.get("math_align", "center")
    set_math_alignment(doc, math_align)

    # Save into an in-memory buffer instead of back to disk.
    # This avoids holding a file handle open on Windows, which causes
    # PermissionError (WinError 32) when TemporaryDirectory tries to clean up.
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



# ---------------------------------------------------------------------------
# /convert  –  POST endpoint
# ---------------------------------------------------------------------------

@app.route("/convert", methods=["POST"])
def convert():
    data = request.get_json(force=True)
    markdown_text = data.get("markdown", "")
    settings = data.get("settings", {})

    # Fix broken table rows caused by newlines around <br> tags
    markdown_text = re.sub(r'(<br\s*/?>)\s*[\r\n]+', r'\1 ', markdown_text, flags=re.IGNORECASE)
    markdown_text = re.sub(r'[\r\n]+\s*(<br\s*/?>)', r' \1', markdown_text, flags=re.IGNORECASE)
    
    # Pandoc ignores <br> inside tables, so we use a safe placeholder to replace later
    markdown_text = re.sub(r'<br\s*/?>', 'XYZZYBRXYZZY', markdown_text, flags=re.IGNORECASE)

    # ম্যাজিক ফিক্স: $$ ব্লকের চারপাশে \n\n দেওয়া হচ্ছে — অংক সুরক্ষিত থাকবে
    markdown_text = re.sub(r'(\$\$[\s\S]*?\$\$)', r'\n\n\1\n\n', markdown_text)

    # ১. অংকের ভেতরের স্পেস ফিক্স 
    markdown_text = re.sub(r'\$(.*?)\$', lambda m: '$' + m.group(1).strip() + '$', markdown_text)

    # ম্যাজিক ফিক্স: $\square$ কে প্লেইন টেক্সট বক্সে রূপান্তর (Word-এ হেডিং হাইড হওয়া আটকাতে)
    markdown_text = re.sub(r'\$\\+square\$', '☐', markdown_text)
    markdown_text = re.sub(r'\$\\+box\$', '☐', markdown_text)
    markdown_text = re.sub(r'\$\s*\\+square\s*\$', '☐', markdown_text)

    # ২. দাঁড়ি (।) ফিক্স - সবচেয়ে নিরাপদ পদ্ধতি (Math block বাদে শুধু সাধারণ টেক্সটে কাজ করবে, টেবিল স্কিপ করবে)
    lines = markdown_text.split('\n')
    for idx, line in enumerate(lines):
        trimmed = line.strip()
        # যদি লাইনটি টেবিল রো হয়, তাহলে দাঁড়ি ফিক্স স্কিপ করবে
        if trimmed.startswith('|') or (line.count('|') >= 2 and not trimmed.startswith('#') and not trimmed.startswith('>')):
            continue
            
        parts = line.split('$')
        for i in range(0, len(parts), 2): # জোড় ইনডেক্সগুলো হলো সাধারণ টেক্সট
            # ক. বাংলা লেখার পর থাকা পাইপ (|) কে দাঁড়ি (।) বানাবে
            parts[i] = re.sub(r'([\u0980-\u09FF]\s*)\|', r'\1।', parts[i])
            # খ. | (২) বা | (৪) এরকম মার্কস এর আগের পাইপকে দাঁড়ি বানাবে
            parts[i] = re.sub(r'\|\s*(?=\()', '। ', parts[i])
        lines[idx] = '$'.join(parts)
    markdown_text = '\n'.join(lines)

    # ৩. উইন্ডোজের হিডেন ক্যারেক্টার রিমুভ
    markdown_text = markdown_text.replace('\r', '')

    # ৪. পারফেক্ট লাইন ব্রেক (টেবিল এবং অংক সুরক্ষিত রেখে)
    math_blocks = []
    def math_extract(m):
        math_blocks.append(m.group(1))
        return f"@@MATH_BLOCK_{len(math_blocks)-1}@@"
    
    # Extract ALL math blocks temporarily
    markdown_text = re.sub(r'(\$\$[\s\S]*?\$\$|\\\[[\s\S]*?\\\]|\\\([\s\S]*?\\\)|\$[^$\n]+\$)', math_extract, markdown_text)
    
    lines = markdown_text.split('\n')
    processed_lines = []
    for i, line in enumerate(lines):
        trimmed = line.strip()
        
        # Magic Fix: Un-indent lines manually padded with spaces (mirrors frontend fix)
        # unless they are lists, quotes, or tables.
        if line.lstrip() != line and not re.match(r'^\s*([\*\-\+]\s|\d+\.\s|>|\|)', line):
            line = line.lstrip()
            
        is_table = trimmed.startswith('|') or ('|' in trimmed and not trimmed.startswith('#') and not trimmed.startswith('>'))
        
        next_trimmed = lines[i+1].strip() if i+1 < len(lines) else ''
        next_is_table = next_trimmed.startswith('|') or ('|' in next_trimmed and not next_trimmed.startswith('#') and not next_trimmed.startswith('>'))
        
        processed_lines.append(line)
        if i < len(lines) - 1:
            if is_table and next_is_table:
                processed_lines.append('\n')
            else:
                processed_lines.append('\n\n')
                
    markdown_text = "".join(processed_lines)
    markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
    
    # Restore math blocks
    for i, block in enumerate(math_blocks):
        markdown_text = markdown_text.replace(f"@@MATH_BLOCK_{i}@@", block)

    markdown_text = re.sub(r'^[-_*]{3,}[ \t]*$', 'SINGLE_LINE_PLACEHOLDER', markdown_text, flags=re.MULTILINE)

    with tempfile.TemporaryDirectory() as tmpdir:
        md_path = os.path.join(tmpdir, "input.md")
        docx_path = os.path.join(tmpdir, "output.docx")

        with open(md_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)

        result = subprocess.run(
            [
                "pandoc", md_path, "-o", docx_path,
                "--from=markdown+tex_math_dollars+tex_math_single_backslash+raw_tex",
                "--standalone",
            ],
            capture_output=True, text=True, timeout=60,
        )

        if result.returncode != 0:
            return jsonify({"error": "Pandoc failed", "details": result.stderr}), 500

        try:
            docx_buffer = apply_docx_settings(docx_path, settings)
        except Exception as exc:
            print(f"[ERROR] Docx styling error: {exc}")
            with open(docx_path, "rb") as f:
                docx_buffer = io.BytesIO(f.read())

    docx_buffer.seek(0)
    return send_file(
        docx_buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="AI2Word_Document.docx",
    )




# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
