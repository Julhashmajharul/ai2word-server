import docx
from docx.oxml.ns import qn
from docx.oxml import parse_xml

doc = docx.Document()
# Add a table, then a paragraph inside a table with math
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0,0)
p = cell.paragraphs[0]
math_xml = '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMathParaPr><m:jc m:val="centerGroup"/></m:oMathParaPr><m:oMath><m:r><m:t>x=y</m:t></m:r></m:oMath></m:oMathPara>'
p._p.append(parse_xml(math_xml))

count_paragraphs = 0
for paragraph in doc.paragraphs:
    count_paragraphs += len(paragraph._element.findall('.//' + qn('m:oMathPara')))

count_body = len(doc._element.body.findall('.//' + qn('m:oMathPara')))

print(f"Found via doc.paragraphs: {count_paragraphs}")
print(f"Found via doc._element.body: {count_body}")
